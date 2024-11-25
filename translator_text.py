import requests
import os
from docx import Document

subscription_Key = ""
endpoint = 'https://api.cognitive.microsofttranslator.com/'
location = "eastus2"
language_destination = 'pt-br'

def translator_text(text, target_language):
    path = '/translate'
    constructed_url = endpoint + path
    headers = {
        'Ocp-Apim-Subscription-Key': subscription_Key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))
    }

    body = [{
        'text': text
    }]

    params = {
        'api-version': '3.0',
        'from': 'en',
        'to': target_language
    }

    try:
        request = requests.post(constructed_url, params=params, headers=headers, json=body)
        request.raise_for_status()  
        response = request.json()
        return response[0]['translations'][0]['text']
    except requests.exceptions.RequestException as e:
        print(f"Erro na requisição: {e}")
        return None
def translate_document (path):
  document = Document(path)
  full_text = []
  for paragraph in document.paragraphs:
    translated_text = translator_text(paragraph.text, language_destination)
    full_text.append(translated_text)

  translated_doc = Document()
  for line in full_text:
    print(line)
    translated_doc.add_paragraph(line)
  path_translated = path.replace(".docx", f"_(language_destination).docx") 
  translated_doc.save(path_translated)
 
  return path_translated
  
input_file = "/content/unwrittenMusic.docx"
translate_document(input_file)


         






  

